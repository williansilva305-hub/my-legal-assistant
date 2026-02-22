import io
import os
import re
import time
import html
import base64
import tempfile
from pathlib import Path
from typing import Optional

import streamlit as st
import streamlit.components.v1 as components
from google import genai
from google.genai import types
from docx import Document
from openpyxl import load_workbook


# ============================================================
# CONFIG GERAL
# ============================================================
st.set_page_config(
    page_title="Falcon",
    page_icon="🦅",
    layout="wide",
)

# ============================================================
# CONSTANTES VISUAIS / BRANDING
# ============================================================
APP_NAME = "Falcon"
BG_COLOR = "#F7F7F8"
SIDEBAR_BG = "#0F0F0F"
USER_BUBBLE = "#002D62"
ASSISTANT_BUBBLE = "#FFFFFF"
ASSISTANT_BORDER = "#E5E5E5"
GOLD = "#C9A227"  # dourado do Falcon Live

# Modelo EXCLUSIVO pedido pelo usuário
MODEL_NAME = "gemini-2.0-flash"

# Link do Falcon Live (pode sobrescrever via Secrets)
# Se quiser abrir arquivo local estático no Streamlit Cloud, habilite:
# .streamlit/config.toml  -> [server] enableStaticServing = true
LIVE_URL = st.secrets.get("LIVE_URL", "/app/static/falcao_live.html")


# ============================================================
# LOGO (static/falcon_logo.png) com corte automático de transparência
# ============================================================
def get_logo_data_uri() -> Optional[str]:
    logo_path = Path("static/falcon_logo.png")
    if not logo_path.exists():
        return None

    try:
        from PIL import Image  # pillow costuma estar disponível
        img = Image.open(logo_path).convert("RGBA")

        # recorta transparência para o símbolo não ficar minúsculo
        alpha = img.getchannel("A")
        bbox = alpha.getbbox()
        if bbox:
            img = img.crop(bbox)

        # adiciona uma margem leve para não "grudar"
        margin = 12
        w, h = img.size
        canvas = Image.new("RGBA", (w + margin * 2, h + margin * 2), (0, 0, 0, 0))
        canvas.paste(img, (margin, margin))
        img = canvas

        bio = io.BytesIO()
        img.save(bio, format="PNG")
        b64 = base64.b64encode(bio.getvalue()).decode("utf-8")
        return f"data:image/png;base64,{b64}"

    except Exception:
        raw = logo_path.read_bytes()
        b64 = base64.b64encode(raw).decode("utf-8")
        return f"data:image/png;base64,{b64}"


LOGO_URI = get_logo_data_uri()


# ============================================================
# CSS CUSTOM (layout + branding)
# ============================================================
sidebar_logo_html = (
    f'<img src="{LOGO_URI}" class="falcon-logo" />' if LOGO_URI else '<div class="falcon-logo-fallback">🦅</div>'
)

st.markdown(
    f"""
<style>
/* ===== Reset visual do Streamlit ===== */
#MainMenu, footer, header {{
    visibility: hidden;
}}
[data-testid="collapsedControl"] {{
    display: none !important;
}}

html, body, [data-testid="stAppViewContainer"] {{
    background: {BG_COLOR} !important;
}}

.block-container {{
    padding-top: 0.7rem !important;
    padding-bottom: 170px !important; /* espaço para input flutuante */
    max-width: 1400px !important;
}}

/* ===== Sidebar ===== */
[data-testid="stSidebar"] {{
    background: {SIDEBAR_BG} !important;
    border-right: 1px solid rgba(255,255,255,0.05);
}}
[data-testid="stSidebar"] * {{
    color: #F5F5F5 !important;
}}

.sidebar-brand {{
    display: flex;
    align-items: center;
    gap: 10px;
    margin: 8px 0 16px 0;
    padding: 4px 2px;
}}
.falcon-logo {{
    width: 70px;
    height: 70px;
    object-fit: contain;
    display: block;
    filter: drop-shadow(0 2px 8px rgba(0,0,0,0.35));
}}
.falcon-logo-fallback {{
    width: 70px;
    height: 70px;
    border-radius: 14px;
    display: flex;
    align-items: center;
    justify-content: center;
    background: rgba(255,255,255,0.05);
    font-size: 32px;
}}
.sidebar-brand-text {{
    display: flex;
    flex-direction: column;
    line-height: 1.05;
}}
.sidebar-brand-text .title {{
    font-weight: 700;
    font-size: 1rem;
    color: #FFFFFF;
}}
.sidebar-brand-text .sub {{
    font-size: 0.78rem;
    color: #A1A1AA;
    margin-top: 3px;
}}

/* Botão ghost + Nova Conversa */
[data-testid="stSidebar"] .stButton > button {{
    width: 100%;
    justify-content: flex-start;
    border-radius: 12px;
    border: 1px solid rgba(255,255,255,0.18) !important;
    background: transparent !important;
    color: #FFFFFF !important;
    padding: 0.65rem 0.85rem;
    box-shadow: none !important;
}}
[data-testid="stSidebar"] .stButton > button:hover {{
    background: rgba(255,255,255,0.05) !important;
    border-color: rgba(255,255,255,0.30) !important;
}}

/* Popover da sidebar (importar) */
[data-testid="stSidebar"] .stPopover > button {{
    width: 100%;
    justify-content: flex-start;
    border-radius: 12px;
    border: 1px solid rgba(255,255,255,0.10) !important;
    background: rgba(255,255,255,0.03) !important;
    color: #FFFFFF !important;
    padding: 0.65rem 0.85rem;
}}
[data-testid="stSidebar"] .stPopover > button:hover {{
    background: rgba(255,255,255,0.06) !important;
}}

/* ===== Header simples ===== */
.falcon-header {{
    display: flex;
    align-items: center;
    justify-content: space-between;
    margin-bottom: 8px;
}}
.falcon-header-left {{
    display: flex;
    flex-direction: column;
}}
.falcon-header-title {{
    font-size: 1.05rem;
    font-weight: 700;
    color: #111827;
}}
.falcon-header-sub {{
    font-size: 0.86rem;
    color: #6B7280;
    margin-top: 2px;
}}

/* ===== Falcon Live (estilo Gemini Live) ===== */
.falcon-live-wrap {{
    display: flex;
    justify-content: center;
    margin: 8px 0 12px 0;
}}
.falcon-live-btn {{
    display: inline-flex;
    align-items: center;
    gap: 8px;
    padding: 10px 16px;
    border-radius: 999px;
    background: #111111;
    color: #F8FAFC !important;
    text-decoration: none !important;
    border: 1px solid {GOLD};
    box-shadow: 0 0 0 1px rgba(201,162,39,0.12), 0 8px 24px rgba(0,0,0,0.15);
    font-weight: 600;
    font-size: 0.92rem;
}}
.falcon-live-btn:hover {{
    background: #161616;
    color: #FFFFFF !important;
    border-color: #E0B93C;
}}
.falcon-live-waves {{
    display: inline-block;
    color: {GOLD};
    font-weight: 700;
    letter-spacing: 1px;
}}

/* ===== Card principal do chat ===== */
.chat-panel {{
    background: #FFFFFF;
    border: 1px solid {ASSISTANT_BORDER};
    border-radius: 16px;
    box-shadow: 0 6px 24px rgba(15,23,42,0.05);
    overflow: hidden;
}}
.chat-scroll {{
    height: calc(100vh - 285px);
    min-height: 340px;
    max-height: calc(100vh - 285px);
    overflow-y: auto;
    padding: 14px;
    background: #FFFFFF;
    scroll-behavior: smooth;
}}

/* ===== Mensagens ===== */
.msg-row {{
    display: flex;
    width: 100%;
    margin: 10px 0;
}}
.msg-row.user {{
    justify-content: flex-end;
}}
.msg-row.assistant {{
    justify-content: flex-start;
}}

.bubble {{
    max-width: 78%;
    border-radius: 18px;
    padding: 10px 12px;
    font-size: 0.96rem;
    line-height: 1.45;
    white-space: pre-wrap;
    word-wrap: break-word;
}}
.bubble.user {{
    background: {USER_BUBBLE};
    color: #FFFFFF;
    border: 1px solid rgba(255,255,255,0.06);
}}
.bubble.assistant {{
    background: {ASSISTANT_BUBBLE};
    color: #111827;
    border: 1px solid {ASSISTANT_BORDER};
    box-shadow: 0 2px 10px rgba(0,0,0,0.04);
}}
.bubble .label {{
    font-weight: 700;
    margin-bottom: 4px;
    color: #111827;
}}

.empty-state {{
    height: 100%;
    min-height: 240px;
    display: flex;
    align-items: center;
    justify-content: center;
    color: #6B7280;
    text-align: center;
    line-height: 1.5;
    font-size: 0.95rem;
}}

.chips-wrap {{
    display: flex;
    flex-wrap: wrap;
    gap: 6px;
    margin: 8px 0 0 0;
}}
.chip {{
    border-radius: 999px;
    border: 1px solid #E5E7EB;
    background: #FFFFFF;
    color: #334155;
    padding: 4px 10px;
    font-size: 0.80rem;
}}

/* ===== Input flutuante (st.chat_input) ===== */
[data-testid="stChatInput"] {{
    position: fixed !important;
    left: 50% !important;
    transform: translateX(-50%) !important;
    bottom: 16px !important;
    width: min(960px, calc(100vw - 330px)) !important; /* espaço da sidebar */
    z-index: 999 !important;
    background: transparent !important;
    padding: 0 !important;
}}
[data-testid="stChatInput"] > div {{
    background: transparent !important;
}}
[data-testid="stChatInput"] textarea {{
    border-radius: 999px !important;   /* pill */
    border: 1px solid #E5E7EB !important;
    background: #FFFFFF !important;
    box-shadow: 0 8px 30px rgba(15,23,42,0.08) !important;
    padding-left: 14px !important;
    padding-right: 52px !important;
    min-height: 46px !important;
}}
[data-testid="stChatInput"] button {{
    border-radius: 999px !important;
    width: 36px !important;
    height: 36px !important;
    min-height: 36px !important;
    margin-right: 6px !important;
}}

/* ===== Ajuste mobile ===== */
@media (max-width: 900px) {{
    .block-container {{
        padding-bottom: 190px !important;
    }}
    .chat-scroll {{
        height: calc(100vh - 340px);
        max-height: calc(100vh - 340px);
    }}
    [data-testid="stChatInput"] {{
        width: calc(100vw - 24px) !important;
        left: 12px !important;
        transform: none !important;
        bottom: 12px !important;
    }}
    .bubble {{
        max-width: 90%;
    }}
}}
</style>
""",
    unsafe_allow_html=True,
)


# ============================================================
# CLIENTE GEMINI (novo SDK)
# ============================================================
API_KEY = st.secrets.get("GEMINI_API_KEY")
if not API_KEY:
    st.error("❌ Configure `GEMINI_API_KEY` em Settings → Secrets do Streamlit.")
    st.stop()


@st.cache_resource
def get_client(api_key: str):
    return genai.Client(api_key=api_key)


client = get_client(API_KEY)


# ============================================================
# PERSONA "MESTRE" (system instruction)
# ============================================================
SYSTEM_INSTRUCTION = """
Você é o Mestre Falcon, um assistente jurídico e professor particular.

PERSONALIDADE E ESTILO:
- Fale em português do Brasil.
- Seja didático, encorajador e natural (fluido, sem soar robótico).
- Explique como um professor que guia o raciocínio.
- Quando útil, use exemplos práticos e perguntas curtas para estimular o pensamento.

PRECISÃO JURÍDICA:
- Seja juridicamente preciso.
- Não invente artigos, súmulas, jurisprudências ou prazos.
- Quando houver incerteza, diga com transparência.
- Diferencie explicação educativa de orientação profissional definitiva.

ANÁLISE DE ANEXOS:
- Se receber documentos, imagens, áudios ou vídeos, organize a resposta.
- Estruture, quando fizer sentido, em:
  1) Fatos
  2) Pontos jurídicos
  3) Riscos
  4) Estratégia sugerida
"""


def create_chat():
    return client.chats.create(
        model=MODEL_NAME,
        config=types.GenerateContentConfig(
            system_instruction=SYSTEM_INSTRUCTION,
            temperature=0.6,
            top_p=0.95,
            max_output_tokens=4096,
        ),
    )


# ============================================================
# SESSION STATE
# ============================================================
if "chat" not in st.session_state:
    st.session_state.chat = create_chat()

if "messages" not in st.session_state:
    st.session_state.messages = []

if "uploader_key" not in st.session_state:
    st.session_state.uploader_key = 0

if "show_history" not in st.session_state:
    st.session_state.show_history = False


# ============================================================
# UTILITÁRIOS (DOCX/XLSX -> TXT para upload)
# ============================================================
def docx_to_text(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    out = []

    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if txt:
            out.append(txt)

    for table in doc.tables:
        out.append("\n[TABELA]")
        for row in table.rows:
            vals = [(c.text or "").replace("\n", " ").strip() for c in row.cells]
            out.append(" | ".join(vals))

    return "\n".join(out).strip()


def xlsx_to_text(file_bytes: bytes, max_rows_per_sheet: int = 200, max_cols: int = 20) -> str:
    wb = load_workbook(io.BytesIO(file_bytes), data_only=True)
    lines = []

    for ws in wb.worksheets:
        lines.append(f"\n### PLANILHA: {ws.title}")
        count = 0
        for row in ws.iter_rows(values_only=True):
            if count >= max_rows_per_sheet:
                lines.append("[... linhas omitidas ...]")
                break
            vals = ["" if c is None else str(c) for c in row[:max_cols]]
            if any(v.strip() for v in vals):
                lines.append(" | ".join(vals))
                count += 1

    return "\n".join(lines).strip()


def normalize_uploaded_file_to_temp(uploaded_file):
    raw = uploaded_file.getvalue()
    ext = Path(uploaded_file.name).suffix.lower()

    if ext == ".docx":
        txt = docx_to_text(raw) or "[DOCX sem texto extraível]"
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".txt")
        tmp.write(txt.encode("utf-8"))
        tmp.flush()
        tmp.close()
        return tmp.name, f"{uploaded_file.name} (documento)"

    if ext in [".xlsx", ".xlsm"]:
        txt = xlsx_to_text(raw) or "[Planilha sem conteúdo legível]"
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".txt")
        tmp.write(txt.encode("utf-8"))
        tmp.flush()
        tmp.close()
        return tmp.name, f"{uploaded_file.name} (planilha)"

    if ext == ".doc":
        raise ValueError(f"{uploaded_file.name}: formato .doc antigo. Converta para .docx ou PDF.")

    # PDF, imagens, áudio, vídeo, txt etc. vão direto
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=ext if ext else ".bin")
    tmp.write(raw)
    tmp.flush()
    tmp.close()
    return tmp.name, uploaded_file.name


def wait_file_active(file_obj, timeout_sec: int = 240):
    start = time.time()
    while True:
        f = client.files.get(name=file_obj.name)
        state = getattr(f, "state", None)
        state_name = getattr(state, "name", str(state)).upper()

        if "ACTIVE" in state_name:
            return f
        if "FAILED" in state_name:
            raise RuntimeError("Falha ao processar um anexo no Gemini.")
        if time.time() - start > timeout_sec:
            raise TimeoutError("O processamento de anexos demorou demais.")
        time.sleep(2)


def upload_attachments(files):
    refs = []
    labels = []
    temp_paths = []

    try:
        for uf in files:
            temp_path, label = normalize_uploaded_file_to_temp(uf)
            temp_paths.append(temp_path)

            uploaded = client.files.upload(file=temp_path)
            uploaded = wait_file_active(uploaded)

            refs.append(uploaded)
            labels.append(label)

        return refs, labels

    finally:
        for p in temp_paths:
            try:
                os.remove(p)
            except Exception:
                pass


# ============================================================
# RENDER HTML DO CHAT (scroll interno)
# ============================================================
def build_chat_html(messages, partial_assistant_text=None):
    rows = []

    if not messages and not partial_assistant_text:
        rows.append(
            '<div class="empty-state">'
            'Pronto para analisar peças, documentos, imagens, áudios e vídeos.<br>'
            'Digite sua mensagem abaixo ou envie anexos pela barra lateral.'
            '</div>'
        )
    else:
        for msg in messages:
            role = msg["role"]
            content = msg["content"]
            safe = html.escape(content).replace("\n", "<br>")
            label = '<div class="label">Mestre:</div>' if role == "assistant" else ""

            rows.append(
                f'<div class="msg-row {role}">'
                f'  <div class="bubble {role}">{label}{safe}</div>'
                f'</div>'
            )

        if partial_assistant_text:
            safe = html.escape(partial_assistant_text).replace("\n", "<br>")
            rows.append(
                '<div class="msg-row assistant">'
                '  <div class="bubble assistant">'
                '    <div class="label">Mestre:</div>'
                f'    {safe}'
                '  </div>'
                '</div>'
            )

    return (
        '<div class="chat-panel">'
        '<div id="falcon-chat-scroll" class="chat-scroll">'
        + "".join(rows)
        + '<div id="falcon-chat-bottom"></div>'
        '</div>'
        '</div>'
    )


def render_chat(chat_placeholder, messages, partial_assistant_text=None):
    chat_placeholder.markdown(
        build_chat_html(messages, partial_assistant_text),
        unsafe_allow_html=True,
    )

    # auto-scroll para a última mensagem
    components.html(
        """
<script>
  const d = window.parent.document;
  const el = d.getElementById("falcon-chat-scroll");
  if (el) {
    el.scrollTop = el.scrollHeight;
  }
</script>
""",
        height=0,
    )


# ============================================================
# SIDEBAR
# ============================================================
with st.sidebar:
    st.markdown(
        f"""
<div class="sidebar-brand">
  {sidebar_logo_html}
  <div class="sidebar-brand-text">
    <div class="title">{APP_NAME}</div>
    <div class="sub">Assistente Jurídico</div>
  </div>
</div>
""",
        unsafe_allow_html=True,
    )

    if st.button("➕  Nova Conversa", use_container_width=True):
        st.session_state.chat = create_chat()
        st.session_state.messages = []
        st.session_state.uploader_key += 1
        st.rerun()

    with st.popover("📎  Importar Documentos", use_container_width=True):
        st.caption("Anexe para a próxima mensagem")
        selected_files_sidebar = st.file_uploader(
            "Arquivos",
            accept_multiple_files=True,
            type=[
                "pdf",
                "png", "jpg", "jpeg", "webp", "bmp",
                "mp3", "wav", "m4a", "aac", "ogg", "flac",
                "mp4", "mov", "avi", "webm", "mkv", "mpeg",
                "txt", "md", "csv", "json", "html", "xml", "rtf",
                "docx", "xlsx", "xlsm",
            ],
            label_visibility="collapsed",
            key=f"uploader_{st.session_state.uploader_key}",
        )
        st.caption("PDF, imagens, áudio, vídeo, DOCX, XLSX...")

    if st.button("🕘  Histórico (sessão)", use_container_width=True):
        st.session_state.show_history = not st.session_state.show_history

    if st.session_state.show_history:
        st.markdown("---")
        msgs = [m["content"] for m in st.session_state.messages if m["role"] == "user"]
        if not msgs:
            st.caption("Sem mensagens ainda.")
        else:
            for i, txt in enumerate(msgs[-8:], 1):
                preview = txt.replace("\n", " ")
                if len(preview) > 40:
                    preview = preview[:40] + "..."
                st.caption(f"{i}. {preview}")

selected_files = locals().get("selected_files_sidebar") or []


# ============================================================
# HEADER + FALCON LIVE
# ============================================================
st.markdown(
    f"""
<div class="falcon-header">
  <div class="falcon-header-left">
    <div class="falcon-header-title">{APP_NAME}</div>
    <div class="falcon-header-sub">Mestre jurídico com respostas fluidas e análise de anexos.</div>
  </div>
</div>
""",
    unsafe_allow_html=True,
)

st.markdown(
    f"""
<div class="falcon-live-wrap">
  <a class="falcon-live-btn" href="{html.escape(LIVE_URL)}" target="_blank">
    🎧 Falcon Live
    <span class="falcon-live-waves">〰〰〰</span>
  </a>
</div>
""",
    unsafe_allow_html=True,
)


# ============================================================
# CHAT (render)
# ============================================================
chat_placeholder = st.empty()
render_chat(chat_placeholder, st.session_state.messages)

if selected_files:
    chips = "".join([f'<span class="chip">📎 {html.escape(f.name)}</span>' for f in selected_files])
    st.markdown(f'<div class="chips-wrap">{chips}</div>', unsafe_allow_html=True)


# ============================================================
# INPUT (st.chat_input obrigatório)
# ============================================================
pergunta = st.chat_input("Digite sua mensagem...")


# ============================================================
# STREAMING palavra por palavra + cursor ▌
# ============================================================
def stream_word_by_word(chat_obj, payload):
    partial = ""

    # MANDATÓRIO: send_message_stream
    for chunk in chat_obj.send_message_stream(payload):
        txt = getattr(chunk, "text", None)
        if not txt:
            continue

        # Divide em palavras + espaços, para efeito "word-by-word"
        parts = re.split(r"(\s+)", txt)
        for part in parts:
            if part == "":
                continue
            partial += part
            render_chat(chat_placeholder, st.session_state.messages, partial_assistant_text=partial + "▌")

    return partial.strip()


def friendly_error_message(exc: Exception) -> str:
    raw = str(exc)

    # quota estourada / resource exhausted
    if "RESOURCE_EXHAUSTED" in raw or "Quota exceeded" in raw or "429" in raw:
        retry_match = re.search(r"retryDelay['\"]?:\s*['\"]?([0-9]+s)", raw)
        retry_text = retry_match.group(1) if retry_match else None

        if retry_text:
            return (
                f"Limite gratuito da API do Gemini atingido no momento. "
                f"Tenta novamente em cerca de {retry_text}. "
                f"(modelo: {MODEL_NAME})"
            )
        return (
            f"Limite gratuito da API do Gemini atingido no momento. "
            f"Tenta novamente em instantes. (modelo: {MODEL_NAME})"
        )

    # modelo não encontrado
    if "NotFound" in raw or "404" in raw:
        return (
            f"Não encontrei o modelo `{MODEL_NAME}` para esta chave/API. "
            f"Confere se a Gemini API está ativa na tua conta e se o projeto está correto."
        )

    # fallback genérico (sem dump gigante)
    return f"Erro ao gerar resposta: {raw[:350]}"


# ============================================================
# ENVIO
# ============================================================
if pergunta and pergunta.strip():
    pergunta = pergunta.strip()

    user_display = pergunta
    if selected_files:
        user_display += "\n\n📎 Anexos enviados: " + ", ".join([f.name for f in selected_files])

    # adiciona msg usuário
    st.session_state.messages.append({"role": "user", "content": user_display})
    render_chat(chat_placeholder, st.session_state.messages)

    try:
        refs = []
        if selected_files:
            with st.spinner("🦅 Processando anexos..."):
                refs, labels = upload_attachments(selected_files)

            st.session_state.messages.append({
                "role": "assistant",
                "content": "✅ Anexos processados: " + " • ".join(labels)
            })
            render_chat(chat_placeholder, st.session_state.messages)

        payload = [*refs, pergunta] if refs else pergunta

        final_text = stream_word_by_word(st.session_state.chat, payload)
        if not final_text:
            final_text = "Não consegui responder agora. Tenta reformular a pergunta."

        # fixa resposta final sem cursor
        st.session_state.messages.append({"role": "assistant", "content": final_text})

        # limpa uploader (anexos ficam prontos só para 1 envio)
        st.session_state.uploader_key += 1

        render_chat(chat_placeholder, st.session_state.messages)
        st.rerun()

    except Exception as e:
        msg = friendly_error_message(e)
        st.session_state.messages.append({"role": "assistant", "content": f"⚠️ {msg}"})
        render_chat(chat_placeholder, st.session_state.messages)
